import os
import json
import pandas as pd
from flask import Flask, request, jsonify, send_file, render_template, send_from_directory
from werkzeug.utils import secure_filename
from docx import Document
import requests
from bs4 import BeautifulSoup
from openai import OpenAI
import io
import uuid

app = Flask(__name__, template_folder='templates', static_folder='static')

# --- CONFIGURATION ---
# 填入你的阿里云 DashScope API Key (如果不想使用环境变量)
DASHSCOPE_API_KEY = "sk-730f9c867cd54bf794dfb086253c6e04 " 
# ---------------------

# Configuration
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
CONFIG_FILE = 'config.json'
DEFAULT_CONFIG = {
    "model": "qwen-plus",
    "system_prompt": """你是一个专业的软件测试工程师。你的任务是根据用户提供的产品文档或网页内容，生成标准格式的测试用例。
请仔细阅读内容，提取功能点，并生成包含以下字段的测试用例：
1. 用例编号 (ID)
2. 模块 (Module)
3. 测试项 (Test Item)
4. 前置条件 (Pre-condition)
5. 测试步骤 (Steps)
6. 预期结果 (Expected Result)
7. 优先级 (Priority)

请以JSON格式输出，格式如下：
[
  {
    "ID": "TC001",
    "Module": "Login",
    "Test Item": "Valid Login",
    "Pre-condition": "User is on login page",
    "Steps": "1. Enter valid username\\n2. Enter valid password\\n3. Click Login",
    "Expected Result": "Login successful, redirect to dashboard",
    "Priority": "High"
  }
]
只输出JSON数组，不要包含Markdown代码块标记（```json ... ```）。"""
}

def load_config():
    if os.path.exists(CONFIG_FILE):
        try:
            with open(CONFIG_FILE, 'r', encoding='utf-8') as f:
                return json.load(f)
        except:
            pass
    return DEFAULT_CONFIG.copy()

def save_config(config):
    with open(CONFIG_FILE, 'w', encoding='utf-8') as f:
        json.dump(config, f, ensure_ascii=False, indent=2)

# Helper Functions for Extraction
def extract_from_docx(filepath):
    doc = Document(filepath)
    content = []
    
    # Iterate through elements to maintain some order (simplified approach: all text then all tables)
    # A better approach for context is tricky without low-level xml parsing, 
    # but usually appending tables after text or text-tables-text is enough for LLM.
    
    content.append("--- DOCUMENT TEXT ---")
    for para in doc.paragraphs:
        if para.text.strip():
            content.append(para.text)
            
    content.append("\n--- DOCUMENT TABLES ---")
    for table in doc.tables:
        content.append("\n[Table Start]")
        # Extract headers
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        content.append(" | ".join(headers))
        content.append("-" * (len(headers) * 10))
        
        # Extract rows
        for row in table.rows[1:]:
            row_data = [cell.text.strip() for cell in row.cells]
            content.append(" | ".join(row_data))
        content.append("[Table End]\n")
        
    return "\n".join(content)

def extract_from_url(url):
    try:
        resp = requests.get(url, timeout=10)
        resp.raise_for_status()
        soup = BeautifulSoup(resp.content, 'html.parser')
        
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.extract()
            
        text = soup.get_text(separator='\n')
        # Clean up whitespace
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = '\n'.join(chunk for chunk in chunks if chunk)
        return text
    except Exception as e:
        raise Exception(f"Failed to fetch URL: {str(e)}")

def extract_from_text(filepath):
    with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
        return f.read()

# AI Generation
def generate_test_cases_ai(content):
    config = load_config()
    # 优先使用全局变量，如果没有设置则尝试环境变量
    api_key = DASHSCOPE_API_KEY or os.getenv("DASHSCOPE_API_KEY")
    if not api_key:
        raise Exception("DASHSCOPE_API_KEY environment variable is not set.")
    
    # 去除可能存在的首尾空格，避免 "Illegal header value" 错误
    api_key = api_key.strip()

    client = OpenAI(
        api_key=api_key,
        base_url="https://dashscope.aliyuncs.com/compatible-mode/v1",
    )

    completion = client.chat.completions.create(
        model=config.get("model", "qwen-plus"),
        messages=[
            {'role': 'system', 'content': config.get("system_prompt")},
            {'role': 'user', 'content': f"Here is the product documentation/content:\n\n{content}"}
        ]
    )
    
    response_content = completion.choices[0].message.content
    
    # Clean up if the AI wraps in code blocks despite instructions
    if response_content.startswith("```json"):
        response_content = response_content[7:]
    if response_content.startswith("```"):
        response_content = response_content[3:]
    if response_content.endswith("```"):
        response_content = response_content[:-3]
        
    return json.loads(response_content.strip())

# Routes
@app.route('/')
def index():
    return render_template('index.html')

@app.route('/api/extract', methods=['POST'])
def extract():
    try:
        content = ""
        inputType = request.form.get('inputType')
        
        if inputType == 'url':
            url = request.form.get('url')
            if not url:
                return jsonify({"error": "URL is required"}), 400
            content = extract_from_url(url)
            
        elif inputType == 'file':
            if 'file' not in request.files:
                return jsonify({"error": "No file uploaded"}), 400
            file = request.files['file']
            if file.filename == '':
                return jsonify({"error": "No selected file"}), 400
                
            filename = secure_filename(file.filename)
            filepath = os.path.join(UPLOAD_FOLDER, filename)
            file.save(filepath)
            
            ext = os.path.splitext(filename)[1].lower()
            if ext == '.docx':
                content = extract_from_docx(filepath)
            elif ext in ['.md', '.txt']:
                content = extract_from_text(filepath)
            else:
                return jsonify({"error": "Unsupported file type"}), 400
                
            # Cleanup uploaded file
            os.remove(filepath)
            
        else:
            return jsonify({"error": "Invalid input type"}), 400

        return jsonify({
            "success": True,
            "content": content
        })

    except Exception as e:
        import traceback
        traceback.print_exc()
        print(f"Error: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/generate', methods=['POST'])
def generate():
    try:
        content = request.json.get('content')
        if not content:
            return jsonify({"error": "No content provided"}), 400

        # Call AI
        test_cases = generate_test_cases_ai(content)
        print(test_cases)
        
        # Generate Excel
        df = pd.DataFrame(test_cases)
        output_filename = f"test_cases_{uuid.uuid4().hex}.xlsx"
        output_path = os.path.join(UPLOAD_FOLDER, output_filename)
        df.to_excel(output_path, index=False)
        
        return jsonify({
            "success": True,
            "download_url": f"/api/download/{output_filename}",
            "preview": test_cases[:5] # Send a preview
        })

    except Exception as e:
        import traceback
        traceback.print_exc()
        print(f"Error: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/download/<filename>')
def download_file(filename):
    return send_from_directory(UPLOAD_FOLDER, filename, as_attachment=True)

@app.route('/api/admin/login', methods=['POST'])
def admin_login():
    data = request.json
    if data.get('password') == '123456':
        return jsonify({"success": True})
    return jsonify({"success": False}), 401

@app.route('/api/config', methods=['GET', 'POST'])
def handle_config():
    # In a real app, we should check for session/auth here, 
    # but for this tool we rely on the frontend 'password' gate for simplicity 
    # (or we could require the password in headers for every request).
    # Let's require password in header for safety.
    password = request.headers.get('X-Admin-Password')
    if password != '123456':
        return jsonify({"error": "Unauthorized"}), 401

    if request.method == 'GET':
        return jsonify(load_config())
    
    if request.method == 'POST':
        new_config = request.json
        save_config(new_config)
        return jsonify({"success": True})

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)
